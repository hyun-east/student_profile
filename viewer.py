import sys
import os
import json
import pandas as pd
from PyQt5.QtWidgets import (QApplication, QMainWindow, QTabWidget, QWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QTextEdit, QComboBox, QPushButton, QFileDialog, QScrollArea, QMenuBar, QAction,
                             QMessageBox, QSizePolicy, QColorDialog)
from PyQt5.QtGui import QPixmap, QFont, QTextOption, QFontDatabase, QTextCursor, QColor
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Inches


class StudentEvaluationWidget(QWidget):
    def __init__(self, directory, initial_font):
        super().__init__()
        self.directory = directory
        self.current_font = initial_font

        main_layout = QHBoxLayout()

        # Left panel for tabs of text information
        self.left_panel = QVBoxLayout()
        self.left_tabs = QTabWidget()
        self.create_text_tabs()
        self.left_panel.addWidget(self.left_tabs, 8)

        # Notes section
        self.notes = QTextEdit()
        self.notes.setFont(self.current_font)
        self.left_panel.addWidget(self.notes, 2)

        left_scroll = QScrollArea()
        left_scroll.setWidgetResizable(True)
        left_container = QWidget()
        left_container.setLayout(self.left_panel)
        left_scroll.setWidget(left_container)
        main_layout.addWidget(left_scroll, 4)

        # Right panel for tabs, notes, and evaluation
        right_panel = QVBoxLayout()

        # Tabs for graphs
        self.graph_tabs = QTabWidget()
        self.create_graph_tabs()
        right_panel.addWidget(self.graph_tabs, 4)

        # Evaluation criteria with vertical slider
        self.evaluation_layout = QVBoxLayout()
        self.evaluation_container = QWidget()
        self.evaluation_container.setLayout(self.evaluation_layout)
        evaluation_scroll = QScrollArea()
        evaluation_scroll.setWidgetResizable(True)
        evaluation_scroll.setWidget(self.evaluation_container)
        right_panel.addWidget(evaluation_scroll, 5)

        self.academic_criteria = [
            ("학업 성취도", 15),
            ("학업 태도", 10),
            ("탐구력", 15)
        ]
        self.career_criteria = [
            ("전공관련 교과 이수 노력", 10),
            ("전공관련 교과 성취도", 10),
            ("진로 탐색 활동과 경험", 20)
        ]
        self.community_criteria = [
            ("협업 및 소통 능력", 5),
            ("나눔과 배려", 5),
            ("성실성과 규칙 준수", 5),
            ("리더십", 5)
        ]

        self.evaluation_criteria = []

        self.add_evaluation_section("학업 역량", self.academic_criteria)
        self.add_evaluation_section("진로 역량", self.career_criteria)
        self.add_evaluation_section("공동체 역량", self.community_criteria)

        self.total_score_label = QLabel("총점: 0점")
        self.total_score_label.setFont(self.current_font)
        right_panel.addWidget(self.total_score_label)

        self.save_button = QPushButton('Save')
        self.save_button.setFont(self.current_font)
        self.save_button.clicked.connect(lambda: self.save_notes(show_message=True))
        right_panel.addWidget(self.save_button)

        self.export_button = QPushButton('Export')
        self.export_button.setFont(self.current_font)
        self.export_button.clicked.connect(self.export_to_word)
        right_panel.addWidget(self.export_button)

        main_layout.addLayout(right_panel, 3)
        self.setLayout(main_layout)

        self.load_notes()

    def create_text_tabs(self):
        text_files = {
            '자율': 'jayul.xlsx',
            '진로': 'jinro.xlsx',
            '동아리': 'club.xlsx',
            '세특': 'seteuk.xlsx',
            '행특': 'behave.xlsx'
        }

        for name, file_name in text_files.items():
            file_path = os.path.join(self.directory, file_name)
            tab = QWidget()
            layout = QVBoxLayout()
            try:
                df = pd.read_excel(file_path)
                df = df[['grade', 'contents']]  # Only display 'grade' and 'contents' columns

                contents_with_spacing = '\n\n'.join(df.apply(lambda row: f"{row['grade']}\n{row['contents']}", axis=1))

                data_edit = QTextEdit()
                data_edit.setReadOnly(True)
                data_edit.setFont(self.current_font)
                data_edit.setPlainText(contents_with_spacing)
                data_edit.setWordWrapMode(QTextOption.WordWrap)
                data_edit.setAlignment(Qt.AlignJustify)
                data_edit.setContextMenuPolicy(Qt.CustomContextMenu)
                data_edit.customContextMenuRequested.connect(self.show_context_menu)

                # 기본 하이라이트 색상을 흰색으로 설정
                cursor = data_edit.textCursor()
                cursor.select(QTextCursor.Document)
                fmt = cursor.charFormat()
                fmt.setBackground(QColor(Qt.white))
                cursor.setCharFormat(fmt)

                layout.addWidget(data_edit)
            except Exception as e:
                error_label = QLabel(f"Failed to load {name}: {str(e)}")
                layout.addWidget(error_label)

            tab.setLayout(layout)
            self.left_tabs.addTab(tab, name)

    def show_context_menu(self, position):
        text_edit = self.sender()
        context_menu = text_edit.createStandardContextMenu()
        highlight_action = QAction("하이라이트", self)
        highlight_action.triggered.connect(lambda: self.highlight_text(text_edit))
        context_menu.addAction(highlight_action)
        context_menu.exec_(text_edit.viewport().mapToGlobal(position))

    def highlight_text(self, text_edit):
        cursor = text_edit.textCursor()
        if cursor.hasSelection():
            color = QColorDialog.getColor(Qt.yellow, self, "Select Highlight Color")
            if color.isValid():
                fmt = cursor.charFormat()
                fmt.setBackground(color)
                cursor.setCharFormat(fmt)
                self.save_notes()

    def add_evaluation_section(self, section_name, criteria):
        section_label = QLabel(f"{section_name}")
        section_label.setFont(QFont(self.current_font.family(), self.current_font.pointSize(), QFont.Bold))
        self.evaluation_layout.addWidget(section_label)

        for criterion_name, max_score in criteria:
            layout = QHBoxLayout()
            score_label = QLabel(f"{criterion_name}: 0점 / {max_score}점")
            score_label.setFont(self.current_font)
            combo_box = QComboBox()
            combo_box.addItems([str(i) for i in range(max_score + 1)])
            combo_box.currentIndexChanged.connect(
                self.create_combobox_callback(score_label, criterion_name, max_score, section_name, combo_box))
            layout.addWidget(score_label)
            layout.addWidget(combo_box)
            self.evaluation_layout.addLayout(layout)
            self.evaluation_criteria.append((score_label, combo_box))

    def create_combobox_callback(self, score_label, criterion_name, max_score, section, combo_box):
        def callback(index):
            score = int(combo_box.currentText())
            self.update_score(score_label, score, max_score, section, criterion_name)

        return callback

    def update_score(self, label, score, max_score, section, criterion):
        label.setText(f"{criterion}: {score}점 / {max_score}점")
        self.calculate_total_score()

    def calculate_total_score(self):
        total_score = 0
        section_totals = {"학업 역량": 0, "진로 역량": 0, "공동체 역량": 0}
        for section, criteria in {"학업 역량": self.academic_criteria, "진로 역량": self.career_criteria,
                                  "공동체 역량": self.community_criteria}.items():
            section_score = 0
            for criterion_name, max_score in criteria:
                for label, combo_box in self.evaluation_criteria:
                    if criterion_name in label.text():
                        score = int(combo_box.currentText())
                        section_score += score
                        break
            section_totals[section] = section_score
            total_score += section_score

        self.total_score_label.setText(f"총점: {total_score}점")
        for section, score in section_totals.items():
            section_label = next((lbl for lbl in self.evaluation_layout.children() if
                                  isinstance(lbl, QLabel) and lbl.text().startswith(section)), None)
            if section_label:
                section_label.setText(f"{section} - {score}점")

    def create_graph_tabs(self):
        graphs = {
            '전과목 내신': 'all_naesin.png',
            '국영수': 'korengmath_naesin.png',
            '국영수과': 'science_naesin.png',
            '국영수사': 'social_naesin.png'
        }

        for name, file_name in graphs.items():
            file_path = os.path.join(self.directory, file_name)
            tab = QWidget()
            layout = QVBoxLayout()

            try:
                pixmap = QPixmap(file_path)
                if pixmap.isNull():
                    raise Exception(f"Image file {file_name} could not be loaded.")
                label = QLabel()
                label.setAlignment(Qt.AlignCenter)
                label.setPixmap(pixmap)
                label.setScaledContents(True)  # Ensure the image scales with the label size
                label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
                layout.addWidget(label)
            except Exception as e:
                error_label = QLabel(f"Failed to load {name} graph: {str(e)}")
                layout.addWidget(error_label)

            tab.setLayout(layout)
            self.graph_tabs.addTab(tab, name)

    def load_notes(self):
        notes_file_path = os.path.join(self.directory, 'notes.json')
        if os.path.exists(notes_file_path):
            try:
                with open(notes_file_path, 'r') as file:
                    data = json.load(file)
                    self.notes.setText(data.get('notes', ''))
                    self.load_evaluation_criteria(data.get('evaluation_criteria', {}))
                    self.load_highlights(data.get('highlights', {}))
                    self.set_all_fonts(data.get('font_size', self.current_font.pointSize()))
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load notes: {str(e)}")
        else:
            self.notes.setPlaceholderText("Enter your notes here...")

    def load_evaluation_criteria(self, content):
        try:
            for criterion_name, score in content.items():
                for label, combo_box in self.evaluation_criteria:
                    if criterion_name in label.text():
                        label.setText(f"{criterion_name}: {score}점 / {label.text().split('/')[1]}")
                        combo_box.setCurrentIndex(score)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load evaluation criteria: {str(e)}")

    def load_highlights(self, highlights):
        try:
            for tab_index in range(self.left_tabs.count()):
                tab = self.left_tabs.widget(tab_index)
                text_edit = tab.layout().itemAt(0).widget()
                tab_name = self.left_tabs.tabText(tab_index)
                if tab_name in highlights:
                    for highlight in highlights[tab_name]:
                        cursor = text_edit.textCursor()
                        cursor.setPosition(highlight['start'])
                        cursor.setPosition(highlight['end'], QTextCursor.KeepAnchor)
                        fmt = cursor.charFormat()
                        fmt.setBackground(QColor(highlight['color']))
                        cursor.setCharFormat(fmt)
                else:
                    # 하이라이트 정보가 없는 경우 기본적으로 흰색 하이라이트를 설정합니다.
                    cursor = text_edit.textCursor()
                    cursor.select(QTextCursor.Document)
                    fmt = cursor.charFormat()
                    fmt.setBackground(QColor(Qt.white))
                    cursor.setCharFormat(fmt)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load highlights: {str(e)}")

    def save_notes(self, show_message=False):
        try:
            notes = self.notes.toPlainText()
            evaluation_content = self.generate_evaluation_content()
            highlights = self.generate_highlights()

            data = {
                'notes': notes,
                'evaluation_criteria': evaluation_content,
                'highlights': highlights,
                'font_size': self.current_font.pointSize()
            }

            with open(os.path.join(self.directory, 'notes.json'), 'w') as file:
                json.dump(data, file, ensure_ascii=False, indent=4)

            if show_message:
                QMessageBox.information(self, 'Save', '데이터가 성공적으로 저장되었습니다.')
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save notes: {str(e)}")

    def generate_evaluation_content(self):
        content = {}
        for label, combo_box in self.evaluation_criteria:
            score = int(combo_box.currentText())
            criterion_name = label.text().split(":")[0]
            content[criterion_name] = score
        return content

    def generate_highlights(self):
        highlights = {}
        for tab_index in range(self.left_tabs.count()):
            tab = self.left_tabs.widget(tab_index)
            text_edit = tab.layout().itemAt(0).widget()
            tab_name = self.left_tabs.tabText(tab_index)

            highlights[tab_name] = []
            cursor = text_edit.textCursor()
            cursor.movePosition(QTextCursor.Start)
            while cursor.movePosition(QTextCursor.NextCharacter, QTextCursor.KeepAnchor):
                if cursor.charFormat().background().color() != QColor(Qt.transparent):
                    start = cursor.selectionStart()
                    color = cursor.charFormat().background().color().name()
                    cursor.clearSelection()
                    while cursor.movePosition(QTextCursor.NextCharacter, QTextCursor.KeepAnchor):
                        if cursor.charFormat().background().color().name() != color:
                            break
                    end = cursor.selectionEnd()
                    highlights[tab_name].append({'start': start, 'end': end, 'color': color})
                    cursor.setPosition(end)
        return highlights

    def increase_font_size(self):
        self.current_font.setPointSize(self.current_font.pointSize() + 1)
        self.set_all_fonts(self.current_font.pointSize())

    def decrease_font_size(self):
        if self.current_font.pointSize() > 1:
            self.current_font.setPointSize(self.current_font.pointSize() - 1)
            self.set_all_fonts(self.current_font.pointSize())

    def set_all_fonts(self, font_size):
        self.current_font.setPointSize(font_size)
        self.notes.setFont(self.current_font)
        self.total_score_label.setFont(self.current_font)

        for i in range(self.left_tabs.count()):
            tab = self.left_tabs.widget(i)
            if isinstance(tab, QWidget):
                layout = tab.layout()
                if layout is not None:
                    for j in range(layout.count()):
                        item = layout.itemAt(j).widget()
                        if isinstance(item, QTextEdit):
                            item.setFont(self.current_font)

        for section_label, combo_box in self.evaluation_criteria:
            section_label.setFont(self.current_font)

        for section_label in self.evaluation_layout.children():
            if isinstance(section_label, QLabel):
                section_label.setFont(self.current_font)

        self.save_notes()

    def export_to_word(self):
        try:
            student_name = os.path.basename(self.directory)
            document = Document()

            document.add_heading(f'{student_name} 생기부 분석 레포트', level=1)
            document.add_paragraph('풍산고등학교 19기 김현동')

            document.add_heading('메모 내용', level=2)
            document.add_paragraph(self.notes.toPlainText())

            document.add_heading('평가 점수', level=2)

            section_scores = {
                "학업 역량": 0,
                "진로 역량": 0,
                "공동체 역량": 0
            }

            for section, criteria in [("학업 역량", self.academic_criteria), ("진로 역량", self.career_criteria),
                                      ("공동체 역량", self.community_criteria)]:
                document.add_heading(section, level=3)
                for criterion_name, max_score in criteria:
                    for label, combo_box in self.evaluation_criteria:
                        if criterion_name in label.text():
                            score = int(combo_box.currentText())
                            document.add_paragraph(f"{criterion_name}: {score}점 / {max_score}점")
                            section_scores[section] += score
                            break

            total_score = sum(section_scores.values())
            document.add_heading(f'총점: {total_score}점', level=2)

            document.add_heading('그래프', level=2)
            for name, file_name in [('전과목 내신', 'all_naesin.png'),
                                    ('국영수', 'korengmath_naesin.png'),
                                    ('국영수과', 'science_naesin.png'),
                                    ('국영수사', 'social_naesin.png')]:
                file_path = os.path.join(self.directory, file_name)
                if os.path.exists(file_path):
                    document.add_heading(name, level=3)
                    document.add_picture(file_path, width=Inches(6.0))

            document.add_heading('자율, 진로, 동아리, 세특', level=2)
            for tab_name in ['자율', '진로', '동아리', '세특']:
                for i in range(self.left_tabs.count()):
                    if self.left_tabs.tabText(i) == tab_name:
                        tab = self.left_tabs.widget(i)
                        text_edit = tab.layout().itemAt(0).widget()
                        document.add_heading(tab_name, level=3)
                        document.add_paragraph(text_edit.toPlainText())

            output_path = os.path.join(self.directory, f'{student_name} 생기부 분석 레포트.docx')
            document.save(output_path)
            QMessageBox.information(self, 'Export', '레포트가 성공적으로 저장되었습니다.')
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export to Word: {str(e)}")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle('생활기록부 평가 프로그램 by 풍산고등학교 19기 김현동')
        self.setGeometry(100, 100, 600, 400)

        self.tabs = QTabWidget()
        self.tabs.setTabsClosable(True)
        self.tabs.tabCloseRequested.connect(self.close_tab)
        self.setCentralWidget(self.tabs)

        menubar = self.menuBar()

        # File menu
        file_menu = menubar.addMenu('File')
        add_folder_action = QAction('Open Student Folder', self)
        add_folder_action.triggered.connect(self.add_student_folder)
        file_menu.addAction(add_folder_action)

        # Settings menu
        settings_menu = menubar.addMenu('Settings')
        increase_font_action = QAction('글씨 크기 키우기', self)
        increase_font_action.triggered.connect(self.increase_font_size)
        settings_menu.addAction(increase_font_action)

        decrease_font_action = QAction('글씨 크기 줄이기', self)
        decrease_font_action.triggered.connect(self.decrease_font_size)
        settings_menu.addAction(decrease_font_action)

        self.is_saved = True
        self.font_settings = self.load_font_settings()
        self.current_font = self.load_font()

    def add_student_folder(self):
        directory = QFileDialog.getExistingDirectory(self, "Select Student Folder")
        if directory:
            student_name = os.path.basename(directory)
            student_widget = StudentEvaluationWidget(directory, self.current_font)
            self.tabs.addTab(student_widget, student_name)
            self.tabs.setCurrentWidget(student_widget)
            self.is_saved = False

    def close_tab(self, index):
        widget = self.tabs.widget(index)
        if widget and not self.is_saved:
            response = QMessageBox.question(self, "Unsaved Changes", "변경사항을 저장하시겠습니까? Yes 버튼을 누르면 자동으로 저장됩니다?",
                                            QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)
            if response == QMessageBox.Yes:
                widget.save_notes(show_message=False)
            elif response == QMessageBox.Cancel:
                return
        self.tabs.removeTab(index)
        self.is_saved = True

    def closeEvent(self, event):
        if not self.is_saved:
            response = QMessageBox.question(self, "Unsaved Changes", "변경사항을 저장하시겠습니까? Yes 버튼을 누르면 자동으로 저장됩니다.",
                                            QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)
            if response == QMessageBox.Yes:
                for index in range(self.tabs.count()):
                    widget = self.tabs.widget(index)
                    if widget:
                        widget.save_notes(show_message=False)
                QMessageBox.information(self, 'Save', '데이터가 성공적으로 저장되었습니다.')
                event.accept()
            elif response == QMessageBox.No:
                event.accept()
            else:
                event.ignore()
        else:
            event.accept()

    def increase_font_size(self):
        current_widget = self.tabs.currentWidget()
        if isinstance(current_widget, StudentEvaluationWidget):
            current_widget.increase_font_size()
            self.save_font_size(current_widget.current_font.pointSize())

    def decrease_font_size(self):
        current_widget = self.tabs.currentWidget()
        if isinstance(current_widget, StudentEvaluationWidget):
            current_widget.decrease_font_size()
            self.save_font_size(current_widget.current_font.pointSize())

    def save_font_size(self, font_size):
        self.font_settings['font_size'] = font_size
        self.save_font_settings()

    def load_font_settings(self):
        settings_file = os.path.join('setting', 'settings.json')
        if os.path.exists(settings_file):
            with open(settings_file, 'r') as file:
                return json.load(file)
        return {'font_size': 12}

    def save_font_settings(self):
        settings_file = os.path.join('setting', 'settings.json')
        with open(settings_file, 'w') as file:
            json.dump(self.font_settings, file, ensure_ascii=False, indent=4)

    def load_font(self):
        font_size = self.font_settings['font_size']
        font_path = "setting/NanumGothic.ttf"

        if not os.path.exists(font_path):
            QMessageBox.critical(self, "Error", f"Font file not found: {font_path}")
            return QFont("Arial", font_size)

        font_id = QFontDatabase.addApplicationFont(font_path)
        if font_id == -1:
            QMessageBox.critical(self, "Error", f"Failed to load font: {font_path}")
            return QFont("Arial", font_size)

        loaded_font_families = QFontDatabase.applicationFontFamilies(font_id)
        if not loaded_font_families:
            QMessageBox.critical(self, "Error", f"No font families found in font file: {font_path}")
            return QFont("Arial", font_size)

        loaded_font_family = loaded_font_families[0]
        return QFont(loaded_font_family, font_size)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    main_window = MainWindow()
    main_window.show()
    sys.exit(app.exec_())


